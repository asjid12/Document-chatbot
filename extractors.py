# extractors.py
import os
from typing import List
from pypdf import PdfReader
from docx import Document
from bs4 import BeautifulSoup

def extract_text_from_pdf(file_path: str) -> str:
    """Extracts text from a PDF file."""
    text = ""
    try:
        reader = PdfReader(file_path)
        for page in reader.pages:
            text += page.extract_text() + "\n"
    except Exception as e:
        print(f"Error reading PDF {file_path}: {e}")
        # Depending on requirements, you might want to re-raise or handle differently
        # raise e
        return ""
    return text

def extract_text_from_txt(file_path: str) -> str:
    """Extracts text from a TXT file."""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        print(f"Error reading TXT {file_path}: {e}")
        # raise e
        return ""

def extract_text_from_docx(file_path: str) -> str:
    """Extracts text from a DOCX file."""
    text = ""
    try:
        doc = Document(file_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
        # Optionally handle tables if content is crucial
        # for table in doc.tables:
        #     for row in table.rows:
        #         for cell in row.cells:
        #             text += cell.text + "\t"
        #     text += "\n"
    except Exception as e:
        print(f"Error reading DOCX {file_path}: {e}")
        # raise e
        return ""
    return text

def extract_text_from_html(file_path: str) -> str:
    """Extracts clean text from an HTML file."""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            soup = BeautifulSoup(f, 'html.parser')
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()
        # Get text
        text = soup.get_text()
        # Break into lines and remove leading/trailing space
        lines = (line.strip() for line in text.splitlines())
        # Break multi-headlines into a line each
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        # Drop blank lines
        text = ' '.join(chunk for chunk in chunks if chunk)
        return text
    except Exception as e:
        print(f"Error reading HTML {file_path}: {e}")
        # raise e
        return ""

def extract_text_from_file(file_path: str) -> str:
    """Dispatches to the correct extractor based on file extension."""
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()

    if ext == '.pdf':
        return extract_text_from_pdf(file_path)
    elif ext == '.txt':
        return extract_text_from_txt(file_path)
    elif ext == '.docx':
        return extract_text_from_docx(file_path)
    elif ext == '.html':
        return extract_text_from_html(file_path)
    else:
        raise ValueError(f"Unsupported file type for extraction: {ext}")

# Optional: More sophisticated chunking
# from langchain.text_splitter import RecursiveCharacterTextSplitter
# def chunk_text_advanced(text: str, chunk_size: int = 1000, chunk_overlap: int = 200) -> List[str]:
#     text_splitter = RecursiveCharacterTextSplitter(
#         chunk_size=chunk_size, chunk_overlap=chunk_overlap
#     )
#     docs = text_splitter.create_documents([text])
#     return [doc.page_content for doc in docs]
